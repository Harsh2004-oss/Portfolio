from fastapi import FastAPI, Form, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, StreamingResponse, JSONResponse
from pydantic import BaseModel, EmailStr
from admin import admin_collection, create_default_admin
from db import (
    resume_collection,
    certificate_collection,
    projects_collection,
    summary_collection,
    contact_collection
)
from passlib.hash import bcrypt
from bson import ObjectId
from dotenv import load_dotenv
from datetime import datetime
import os
import docx
import httpx
import requests
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import cloudinary
import cloudinary.uploader

# ==========================
# Load Environment Variables
# ==========================
load_dotenv()
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
EMAIL_USER = os.getenv("EMAIL_USER")
EMAIL_PASS = os.getenv("EMAIL_PASS")
CLOUDINARY_CLOUD_NAME = os.getenv("CLOUDINARY_CLOUD_NAME")
CLOUDINARY_API_KEY = os.getenv("CLOUDINARY_API_KEY")
CLOUDINARY_API_SECRET = os.getenv("CLOUDINARY_API_SECRET")
FRONTEND_URL = os.getenv("FRONTEND_URL")

# ==========================
# Cloudinary Configuration
# ==========================
cloudinary.config(
    cloud_name=CLOUDINARY_CLOUD_NAME,
    api_key=CLOUDINARY_API_KEY,
    api_secret=CLOUDINARY_API_SECRET
)

# ==========================
# Create Default Admin
# ==========================
create_default_admin()

# ==========================
# FastAPI App
# ==========================
app = FastAPI(title="Harsh Portfolio API")

# ==========================
# CORS Configuration
# ==========================
origins = ["http://localhost:5173"]
allowed_origins = os.getenv("ALLOWED_ORIGINS")
if allowed_origins:
    origins.extend([origin.strip() for origin in allowed_origins.split(",")])

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==========================
# Global Exception Handler (ensures CORS headers on 500 errors)
# ==========================
@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    import traceback
    traceback.print_exc()  # Log the full error on Render
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc)},
    )

# ==========================
# ROOT ROUTE
# ==========================
@app.get("/")
def root():
    return {"message": "Portfolio API is running 🚀"}
# ==========================================================
# 🔐 ADMIN LOGIN
# ==========================================================
@app.post("/admin/login")
def login(username: str = Form(...), password: str = Form(...)):
    admin = admin_collection.find_one({"username": username})
    if not admin or not bcrypt.verify(password, admin["password"]):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return {"success": True}

# ==========================================================
# 📄 RESUME — Upload, View, Download
# ==========================================================
@app.post("/admin/upload_resume")
async def upload_resume(file: UploadFile = File(...)):
    if not (file.filename.endswith(".pdf") or file.filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF or DOCX allowed")

    # Read file bytes once
    file_bytes = await file.read()

    # Extract text content for chatbot context
    content = ""
    try:
        if file.filename.lower().endswith(".pdf"):
            import PyPDF2
            from io import BytesIO
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            content = "\n".join([page.extract_text() or "" for page in reader.pages])

        elif file.filename.lower().endswith(".docx"):
            from io import BytesIO
            doc = docx.Document(BytesIO(file_bytes))
            content = "\n".join([para.text for para in doc.paragraphs])

    except Exception as e:
        print(f"⚠️ Text extraction failed: {e}")

    file_base = os.path.splitext(file.filename)[0]

    # Upload using bytes (NOT file.file)
    result = cloudinary.uploader.upload(
        file_bytes,
        resource_type="raw",
        folder="portfolio/resumes",
        public_id=file_base,
        overwrite=True
    )

    resume_collection.delete_many({})
    resume_collection.insert_one({
        "filename": file.filename,
        "file_url": result["secure_url"],
        "content": content
    })

    return {
        "message": "Resume uploaded successfully",
        "file_url": result["secure_url"]
    }



@app.get("/resume")
def get_resume():
    resume = resume_collection.find_one(sort=[("_id", -1)])
    if not resume:
        raise HTTPException(status_code=404, detail="No resume uploaded yet")
    return {
        "filename": resume.get("filename"),
        "file_url": resume.get("file_url"),
        "content": resume.get("content", "")
    }
@app.get("/resume/view")
async def view_resume():
    resume = resume_collection.find_one(sort=[("_id", -1)])
    if not resume or not resume.get("file_url"):
        raise HTTPException(status_code=404, detail="No resume uploaded")

    file_url = resume["file_url"]

    # Fetch file content from Cloudinary
    async with httpx.AsyncClient() as client:
        r = await client.get(file_url)
        if r.status_code != 200:
            raise HTTPException(status_code=404, detail="File not found on Cloudinary")

        return StreamingResponse(
            iter([r.content]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"inline; filename={resume['filename']}"}
        )

# ==========================
# RESUME — Download (as PDF)
# ==========================
@app.get("/resume/download")
async def download_resume():
    resume = resume_collection.find_one(sort=[("_id", -1)])
    if not resume or not resume.get("file_url"):
        raise HTTPException(status_code=404, detail="No resume uploaded")

    file_url = resume["file_url"]

    # Fetch file content from Cloudinary
    async with httpx.AsyncClient() as client:
        r = await client.get(file_url)
        if r.status_code != 200:
            raise HTTPException(status_code=404, detail="File not found on Cloudinary")

        return StreamingResponse(
            iter([r.content]),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={resume['filename']}"}
        )
# ==========================================================
# 📑 SUMMARY UPLOAD
# ==========================================================
@app.post("/admin/upload_summary")
async def upload_summary(file: UploadFile = File(...)):
    if file.filename.endswith(".pdf") or file.filename.endswith(".docx"):
        if file.filename.endswith(".pdf"):
            import PyPDF2
            reader = PyPDF2.PdfReader(file.file)
            content = "\n".join([page.extract_text() or "" for page in reader.pages])
        else:
            doc = docx.Document(file.file)
            content = "\n".join([para.text for para in doc.paragraphs])
    elif file.filename.endswith(".txt"):
        content = (await file.read()).decode("utf-8")
    else:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX or TXT allowed")

    summary_collection.delete_many({})
    summary_collection.insert_one({
        "filename": file.filename,
        "content": content
    })

    return {"message": "Summary uploaded successfully"}

# ==========================================================
# 🏆 CERTIFICATES
# ==========================================================
@app.post("/admin/upload_certificate")
async def upload_certificate(
    title: str = Form(...),
    description: str = Form(...),
    file: UploadFile = File(...)
):
    file_ext = os.path.splitext(file.filename)[1].lower()

    # Use 'raw' for PDFs (so they're accessible), 'auto' for images
    res_type = "raw" if file_ext == ".pdf" else "auto"

    result = cloudinary.uploader.upload(
    file.file,
    resource_type=res_type,
    folder="portfolio/certificates",
    public_id=os.path.splitext(file.filename)[0],
    overwrite=True
)

    certificate_collection.insert_one({
        "title": title,
        "description": description,
        "file_url": result["secure_url"]
    })

    return {"message": "Certificate uploaded successfully", "file_url": result["secure_url"]}


@app.get("/certificates")
def get_certificates():
    certs = list(certificate_collection.find())
    return [
        {
            "id": str(c["_id"]),
            "title": c.get("title"),
            "description": c.get("description"),
            "file_url": c.get("file_url")
        }
        for c in certs
    ]


@app.get("/certificates/view/{cert_id}")
async def view_certificate(cert_id: str):
    cert = certificate_collection.find_one({"_id": ObjectId(cert_id)})
    if not cert or not cert.get("file_url"):
        raise HTTPException(status_code=404, detail="Certificate not found")

    file_url = cert["file_url"]

    # Try to fetch and serve inline
    # Cloudinary may store under different delivery types (image/raw/video)
    # Try the original URL first, then try raw variant for PDFs
    urls_to_try = [file_url]

    # If stored as image/upload, also try raw/upload for PDFs
    if "/image/upload/" in file_url and file_url.lower().endswith(".pdf"):
        urls_to_try.append(file_url.replace("/image/upload/", "/raw/upload/"))

    async with httpx.AsyncClient(timeout=30.0, follow_redirects=True) as client:
        for url in urls_to_try:
            try:
                response = await client.get(url)
                if response.status_code == 200:
                    # Detect content type
                    content_type = response.headers.get("content-type", "application/octet-stream")
                    if url.lower().endswith(".pdf"):
                        content_type = "application/pdf"

                    return StreamingResponse(
                        iter([response.content]),
                        media_type=content_type,
                        headers={"Content-Disposition": "inline"}
                    )
            except Exception:
                continue

    # Fallback: redirect to the original URL directly
    return RedirectResponse(url=file_url)

# ==========================================================
# 🚀 PROJECTS
# ==========================================================
@app.post("/admin/add_project")
async def add_project(
    title: str = Form(...),
    description: str = Form(...),
    tech_stack: str = Form(""),
    github_link: str = Form(""),
    live_link: str = Form(""),
    screenshot: UploadFile = File(None)
):
    image_url = None
    if screenshot and screenshot.filename:
        result = cloudinary.uploader.upload_large(
            screenshot.file,
            resource_type="auto",
            folder="portfolio/projects",
            public_id=os.path.splitext(screenshot.filename)[0]
        )
        image_url = result["secure_url"]

    projects_collection.insert_one({
        "title": title,
        "description": description,
        "tech_stack": tech_stack,
        "github_link": github_link,
        "live_link": live_link,
        "image_url": image_url
    })

    return {"message": "Project added successfully", "image_url": image_url}


@app.put("/admin/update_project/{project_id}")
async def update_project(
    project_id: str,
    title: str = Form(...),
    description: str = Form(...),
    tech_stack: str = Form(""),
    github_link: str = Form(""),
    live_link: str = Form(""),
    screenshot: UploadFile = File(None)
):
    update_data = {
        "title": title,
        "description": description,
        "tech_stack": tech_stack,
        "github_link": github_link,
        "live_link": live_link,
    }

    if screenshot and screenshot.filename:
        result = cloudinary.uploader.upload_large(
            screenshot.file,
            resource_type="auto",
            folder="portfolio/projects",
            public_id=os.path.splitext(screenshot.filename)[0]
        )
        update_data["image_url"] = result["secure_url"]

    result = projects_collection.update_one(
        {"_id": ObjectId(project_id)},
        {"$set": update_data}
    )

    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Project not found")

    return {"message": "Project updated successfully"}


@app.get("/projects")
def get_projects():
    projects = list(projects_collection.find())
    return [
        {
            "id": str(p["_id"]),
            "title": p.get("title"),
            "description": p.get("description"),
            "tech_stack": p.get("tech_stack"),
            "github_link": p.get("github_link"),
            "live_link": p.get("live_link"),
            "image_url": p.get("image_url")
        }
        for p in projects
    ]


@app.delete("/admin/delete_project/{project_id}")
def delete_project(project_id: str):
    projects_collection.delete_one({"_id": ObjectId(project_id)})
    return {"message": "Project deleted successfully"}

# ==========================================================
# 📬 CONTACT FORM
# ==========================================================
class ContactRequest(BaseModel):
    name: str
    email: EmailStr
    subject: str | None = ""
    message: str

@app.post("/contact")
def contact(data: ContactRequest):
    # Save to database first (this is the important part)
    contact_collection.insert_one({
        "name": data.name,
        "email": data.email,
        "subject": data.subject,
        "message": data.message,
        "created_at": datetime.utcnow()
    })

    # Try to send email notification (best-effort, don't fail the request)
    email_sent = False
    if EMAIL_USER and EMAIL_PASS:
        try:
            msg = MIMEMultipart()
            msg["From"] = EMAIL_USER
            msg["To"] = EMAIL_USER
            msg["Subject"] = data.subject or "New Portfolio Contact Message"

            body = f"""
New message from portfolio:

Name: {data.name}
Email: {data.email}
Subject: {data.subject}

Message:
{data.message}
"""
            msg.attach(MIMEText(body, "plain"))

            with smtplib.SMTP("smtp.gmail.com", 587, timeout=10) as server:
                server.starttls()
                server.login(EMAIL_USER, EMAIL_PASS)
                server.send_message(msg)
            email_sent = True
        except Exception as e:
            print(f"⚠️ Email notification failed: {e}")

    return {"message": "Message sent successfully", "email_sent": email_sent}


@app.get("/admin/contacts")
def get_contacts():
    contacts = list(contact_collection.find().sort("created_at", -1))
    return [
        {
            "id": str(c["_id"]),
            "name": c.get("name"),
            "email": c.get("email"),
            "subject": c.get("subject"),
            "message": c.get("message"),
            "created_at": c.get("created_at").isoformat() if c.get("created_at") else None
        }
        for c in contacts
    ]

# ==========================================================
# 🤖 AI CHAT
# ==========================================================
class ChatRequest(BaseModel):
    question: str

@app.post("/chat")
async def chat(data: ChatRequest):
    user_question = data.question.strip().lower()
    greetings = ["hi", "hello", "hey", "good morning", "good afternoon", "good evening", "hii", "hiii", "yo", "sup", "hola", "namaste"]
    if any(user_question.startswith(g) for g in greetings):
        return {"answer": f"{data.question}! I'm Harsh's AI assistant. How can I help you with his skills or projects?"}

    resume = resume_collection.find_one(sort=[("_id", -1)])
    summary = summary_collection.find_one(sort=[("_id", -1)])

    prompt = f"""
You are Harsh's professional AI assistant.
Answer ONLY using the information below.
If answer not found say:
"I'm not sure about that, but I can help with questions related to Harsh's skills, experience, or projects."

RESUME:
{resume.get('content', '') if resume else ''}

SUMMARY:
{summary.get('content', '') if summary else ''}

User Question:
{data.question}
"""

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        },
        json={
            "model": "meta-llama/llama-3-8b-instruct",
            "messages": [{"role": "user", "content": prompt}]
        }
    )

    if response.status_code != 200:
        raise HTTPException(status_code=500, detail="AI API Error")

    answer = response.json()["choices"][0]["message"]["content"]
    return {"answer": answer}